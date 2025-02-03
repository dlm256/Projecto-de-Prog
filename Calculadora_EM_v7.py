import pickle
import numpy as np
from docx import Document
import matplotlib.pyplot as plt
import pandas as pd

df = pd.DataFrame()


Objectos = []
Massa = []
Velocidade = []
EngCin = []
Altura = []
Epg = []
Em = []


def menu_df():
    while True:
        try:
            print("1. Criar tratamento de dados em Data Frame")
            print("2. Consultar tabela de dados")
            print("3. Eliminar coluna da tabela de dados")
            print("4. Exportar tabela para Excel")
            print("5. Importar tabela de Excel")
            print("6. Regressar ao Menu Principal")
            
            opção = int(input("\nIntroduza a opção: " ))
            
            if opção == 1:
                const_DF()
            
            elif opção == 2:
                consultar_DF()
            
            elif opção == 3:
                eliminar_coluna()
                
            elif opção == 4:
                export_xl()
            
            elif opção == 5:
                importar_xl()
                
            elif opção == 6:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")       


def const_DF():
    global df
    if not Objectos:
        print("A lista está vazia!")
    else:
        df = pd.DataFrame({"Objecto": Objectos, "Massa(kg)": Massa, "Velocidade(m/s)": Velocidade, "Altura(m)": Altura, "Energia Cinética(J)": EngCin, "Energia Potencial Gravítica(J)": Epg, "Energia Mecânica(J)":Em})
        print("\nData Frame criado com sucesso!\n") 
        return df

def consultar_DF():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    
    print("Composição do Data Frame: \n")
    print(df)
    print()
    
def eliminar_coluna():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    nome_coluna = input("Indique o nome da Coluna a eliminar: \n") 
    
    if nome_coluna not in df.columns:
        print(f"Erro! A coluna '{nome_coluna}' não consta no Data Frame!\n")
        return None  
    
    df.drop(columns=[nome_coluna], inplace=True)
    print(f"Coluna '{nome_coluna}' eliminada com sucesso!\n")
    print()
    
def export_xl():
    global df
    if df.empty:
        print("Data Frame vazio! ")
        return
    nome_ficheiro = input("Introduza o nome do ficheiro a guardar em excel: ")
    
    if not nome_ficheiro.endswith('.xlsx'):
        nome_ficheiro += '.xlsx'
        
    try:
        df.to_excel(nome_ficheiro, index=False)
        print(f"Dados do Data Frame exportados para: {nome_ficheiro}\n")
    except Exception as e:
        print(f"Erro a gerar ficheiro: {e} \n")
    print()

def importar_xl():
    global df
    
    nome_ficheiro = input("Introduza a localização do ficheiro a importar: \n")
    
    if not nome_ficheiro.endswith('.xlsx'):
        nome_ficheiro += '.xlsx'
        
    try:
        df = pd.read_excel(nome_ficheiro)
        print(f"{nome_ficheiro} importado com sucesso!")
    except FileNotFoundError:
        print("Erro! Ficheiro não encontrado!")
    except Exception as e:
        print(f"Erro a gerar ficheiro: {e} \n")
    print()


def construir_array():
    while True:
        try:
            print("1. Registar dados")
            print("2. Regressar ao Menu Principal")
            opção = int(input("\nIntroduza a opção: " ))
            if opção == 1:
                
                objecto = input("Indique o objecto a adicionar: ")
                while True:
                    try:
                        massa = float(input("Indique a massa em kg: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                while True:
                    try:        
                        velocidade = float(input("Indique a velocidade em m/s: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                while True:
                    try: 
                        altura = float(input("Indique a altura do objecto em relação ao nível do mar em metros: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                calc_ec = 0.5*(massa*(velocidade**2))
                calc_epg = massa*altura*9.8
                calc_em = calc_ec + calc_epg 
                Objectos.append(objecto)
                Massa.append(massa)
                Velocidade.append(velocidade)
                Altura.append(altura)
                EngCin.append(calc_ec)
                Epg.append(calc_epg)
                Em.append(calc_em)
                print("Registado!")
                
            elif opção == 2:
                print("Regressar ao menu principal...")
                break
            else:
                print("Opção inválida! Seleccione novamente: ")
                    
        except ValueError:
            print(" Erro! Indique um número!")
            
            

def calc_media ():
    if not Objectos:
        print("A lista está vazia!")      
    else:                               
        media_massa = np.mean(Massa)
        media_velocidade = np.mean(Velocidade)
        media_altura = np.mean(Altura)
        media_engcin = np.mean(EngCin)
        media_epg = np.mean(Epg)
        media_em = np.mean(Em)
        print(f"Média das massas: {media_massa} kg") 
        print(f"Média das velocidades: {media_velocidade} m/s")
        print(f"Média das alturas: {media_altura} m")
        print(f"Média de Energia Cinética: {media_engcin} joules")
        print(f"Média da Energia potencial Gravítica: {media_epg} joules")
        print(f"Média da Energia Mecânica: {media_em} joules")
        
def somatório():
    if not Objectos:
        print("A lista está vazia!") 
    else:
        sum_massa = sum(Massa)
        sum_vel = sum(Velocidade)
        sum_h = sum(Altura)
        sum_epg = sum(Epg)
        sum_EngCin = sum(EngCin)
        sum_em = sum(Em)
        print(f"Somatório massas: {sum_massa} kg")
        print(f"Somatório velocidades: {sum_vel}")
        print(f"Somatório velocidades: {sum_h}")
        print(f"Somatório Energia Cinética: {sum_EngCin}")
        print(f"Somatório Energia potencial Gravítica: {sum_epg}")
        print(f"Somatório Energia Mecânica: {sum_em}")
        
def maior_menor():
    if not Objectos:
        print("A lista está vazia!") 
    else:
        maior_valor_massa = np.max(Massa)
        menor_val_massa = np.min(Massa)
        print(f"Massa \nMaior valor: {maior_valor_massa} kg \nMenor valor {menor_val_massa} kg")

        maior_valor_vel = np.max(Velocidade)
        menor_val_vel = np.min(Velocidade)
        print(f"Velocidade \nMaior valor: {maior_valor_vel} m/s \nMenor valor {menor_val_vel} m/s")
        
        maior_valor_h = np.max(Altura)
        menor_val_h = np.min(Altura)
        print(f"Altura \nMaior valor: {maior_valor_h} m \nMenor valor {menor_val_h} m")
        
        maior_valor_engcin = np.max(EngCin)
        menor_val_engcin = np.min(EngCin)
        print(f"Energia Cinética \nMaior valor: {maior_valor_engcin} joules \nMenor valor {menor_val_engcin} joules")
        
        maior_valor_epg = np.max(Epg)
        menor_val_epg = np.min(Epg)
        print(f"Energia Potencial Gravítica \nMaior valor: {maior_valor_epg} joules \nMenor valor {menor_val_epg} joules")
        
        maior_valor_em = np.max(Em)
        menor_val_em = np.min(Em)
        print(f"Energia Mecânica \nMaior valor: {maior_valor_em} joules \nMenor valor {menor_val_em} joules")

   
def consultar_array():
    if not Objectos:
        print("A lista está vazia!")
    else:
        print("Lista de Objectos:\n")
        for i, (objecto, massa, velocidade, altura, calc_ec, calc_epg, calc_em) in enumerate(zip(Objectos,Massa,Velocidade,Altura,EngCin,Epg,Em), start=1):
            print(f"Objecto {i}: {objecto}")
            print(f"Massa: {massa} kg")
            print(f"Velocidade: {velocidade} m/s")
            print(f"Altura: {altura} m")
            print(f"Energia Cinética: {calc_ec} joules")
            print(f"Energia Potencial Gravítica: {calc_epg} joules")
            print(f"Energia Mecânica: {calc_em} joules\n")
            
def eliminar_objectos():
    if not Objectos:
        print("A lista está vazia!")
    else:
        consultar_array()
        try:
            while True:
                id = int(input("\nIndique o id do objecto a eliminar: "))
                if 1 <= id <= len(Objectos):
                    objecto_eliminado = Objectos.pop(id -1)
                    massa_eliminada = Massa.pop(id -1)
                    velocidade_eliminada = Velocidade.pop(id -1)
                    h_eliminada = Altura.pop(id -1)
                    ec_eliminada = EngCin.pop(id -1)
                    epg_eliminada = Epg.pop(id -1)
                    em_eliminada = Em.pop(id -1)
                    print(f"\nItem com id {id} eliminado!")
                    break
        except ValueError:
            print("Erro! O id deve ser um numero!")
            
def alterar_objectos():
    if not Objectos:
        print("A lista está vazia!")
    
    else:
        consultar_array()
        try:
            while True:
                id = int(input("\nIndique o id do objecto a alterar: "))
                if 1 <= id <= len(Objectos): 
                    objecto_actualizado = input("Indique um novo objecto: ")
                    massa_actualizada = float(input("Indique uma nova massa: "))
                    h_actualizada = float(input("Indique uma nova altura: "))
                    velocidade_actualizada = float(input("Indique uma nova velocidade: "))
                    
                    Objectos[id -1] = objecto_actualizado
                    Massa[id -1] = massa_actualizada
                    Velocidade [id -1] = velocidade_actualizada
                    Altura [id -1] = h_actualizada
                    EngCin[id -1] = 0.5*(massa_actualizada*(velocidade_actualizada**2))
                    Epg[id -1] = massa_actualizada*h_actualizada*9.8
                    Em[id -1] = EngCin[id -1] + Epg[id -1]
                    print("\nObjecto actualizado com sucesso!")
                    break
                else:
                    print("Erro! O id é inválido!")
        except ValueError:
                print("Erro! O id deve ser um numero!")

def guardar_ficheiro ():
    if not Objectos:
        print("A lista está vazia!")
    else:
        try:
            file_path = input("Introduza a directoria para guardar:  ")
            with open(file_path, "wb") as ficheiro:
                pickle.dump((Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em), ficheiro)
                print("Exportado com sucesso!")
        except Exception as e:
            print(f"Erro ao gerar ficheiro: {e}")


def abrir_ficheiro ():
    global Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em
    
    try:
        file_path = input(" Introduza a directoria para abrir doc: ")
        with open(file_path, "rb") as ficheiro:
            Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em  = pickle.load
            print("Ficheiro importado!")
    except FileNotFoundError:
        print("Erro! Ficheiro não encontrado!")
    except Exception as e:
        print(f"Erro a abrir ficheiro: {e}")
        
def gerar_word():
    if not Objectos:
        print("A lista está vazia!")
    else:
        try:
            doc = Document()
            título = input("Indique titulo: ")
            doc.add_heading(título, 0)

            for i, (objecto, massa, altura, velocidade, calc_ec, calc_epg, calc_em) in enumerate(zip(Objectos,Massa,Altura,Velocidade,EngCin,Epg,Em), start=1):
                doc.add_paragraph(f"{i} - {objecto} com {massa} kg e {velocidade} m/s produz {calc_ec} joules (Energia Cinética), {calc_epg} joules para {altura} metros (Energia Potencial Gravitica) e totalizando {calc_em} joules (Energia Mecânica) \n")
                
            file_path = input("Introduza a directoria para guardar:  ")
            doc.save(file_path)
            print("Doc Word gerado com sucesso!")
        except Exception as e:
            print(f"Erro ao gerar ficheiro: {e}")

def graph():
    while True:
        try:
            print("1. Gráfico de Linha da Relação Velocidade/Energia cinética")
            print("2. Gráfico de Linha da Relação Massa/Energia cinética")
            print("3. Gráfico de Linha da Relação Massa/Energia Potencial Gravítica")
            print("4. Gráfico de Linha da Relação Altura/Energia Potencial Gravítica")
            print("5. Gráfico de Linha da Relação Velocidade/Energia cinética")
            print("6. Gráfico de Linha da Relação Energia Mecânica/Energia Potencial Gravítica")
            print("7. Gráfico de Linha da Relação Energia Mecânica/Energia Cinética")
            print("8. Saír")
            
            opção = int(input("Introduza a opção: " ))
            if opção == 1: 
                plt.plot(Velocidade,EngCin)
                plt.xlabel('Velocidade (m/s)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Velocidade/Energia cinética')
                plt.show()
            elif opção == 2:
                plt.plot(Massa,EngCin)
                plt.xlabel('Massa (kg)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Massa/Energia cinética')
                plt.show()
            elif opção == 3:
                plt.plot(Massa,Epg)
                plt.xlabel('Massa (kg)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Massa/Energia Potencial Gravítica')
                plt.show()
            elif opção == 4:    
                plt.plot(Altura,Epg)
                plt.xlabel('Altura (m)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Altura/Energia Potencial Gravítica')
                plt.show()
            elif opção == 5:   
                plt.plot(EngCin,Epg)
                plt.xlabel('Energia Cinética (joules)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Velocidade/Energia cinética')
                plt.show()
            elif opção == 6:    
                plt.plot(Em,Epg)
                plt.xlabel('Energia Mecânica (joules)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Energia Mecânica/Energia Potencial Gravítica')
                plt.show()
            elif opção == 7:    
                plt.plot(Em,EngCin)
                plt.xlabel('Energia Mecânica (joules)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Energia Mecânica/Energia Cinética')
                plt.show()
            elif opção == 8:
                print("A saír para menu anterior...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
        except ValueError:
                print(" Erro! Indique um número!")
    
def graph_bar():
    while True:
        try:
            print("1. Gráfico de Barras - Velocidade")
            print("2. Gráfico de Barras - Massa")
            print("3. Gráfico de Barras - Altura")
            print("4. Gráfico de Barras - Energia Cinética")
            print("5. Gráfico de Barras - Energia Potencial Gravítica")
            print("6. Gráfico de Barras - Energia Mecânica")
            print("7. Saír")
            opção = int(input("Introduza a opção: " ))
            if opção == 1:
                plt.bar(Objectos,Velocidade)
                plt.xlabel('Objecto')
                plt.ylabel('Velocidade(m/s)')
                plt.title('Gráfico de Barras - Velocidade')
                plt.show()
            elif opção == 2:
                plt.bar(Objectos,Massa)
                plt.xlabel('Objecto')
                plt.ylabel('Massa(kg)')
                plt.title('Gráfico de Barras - Massa')
                plt.show()
            elif opção == 3:  
                plt.bar(Objectos,Altura)
                plt.xlabel('Objecto')
                plt.ylabel('Altura (m)')
                plt.title('Gráfico de Barras - Altura')
                plt.show()
            elif opção == 4:
                plt.bar(Objectos,EngCin)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Cinética(joules)')
                plt.title('Gráfico de Barras - Energia Cinética')
                plt.show()
            elif opção == 5:
                plt.bar(Objectos,Epg)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Potencial Gravítica(joules)')
                plt.title('Gráfico de Barras - Energia Potencial Gravítica')
                plt.show()
            elif opção == 6:
                plt.bar(Objectos,Em)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Mecânica(joules)')
                plt.title('Gráfico de Barras - Energia Mecânica')
                plt.show()
            elif opção == 7:
                print("A saír para Menu anterior...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
            
            
        except ValueError:
                print(" Erro! Indique um número!")
            
def menu_graph():
    while True:
        try:
            print("1. Gráficos de linhas da relação entre os valores")
            print("2. Gráficos de Barras para atribuição de valor por objecto")
            print("3. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                graph()
            
            elif opção == 2:
                graph_bar()
            
            elif opção == 3:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def calc_adicionais():
    while True:
        try:
            print("1. Calcular a média")
            print("2. Calcular o somatório de cada conjunto de dados")
            print("3. Calcular Maior e Menor de cada conjunto")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                calc_media()
            
            elif opção == 2:
                somatório()
                
            elif opção == 3:
                maior_menor()
            
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def gestao_array():
    while True:
        try:
            print("1. Listar objectos")
            print("2. Eliminar objectos")
            print("3. Alterar objectos")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                consultar_array()
            
            elif opção == 2:
                eliminar_objectos()
                
            elif opção == 3:
                alterar_objectos()
            
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def armazenamento():
    while True:
        try:
            print("1. Guardar dados em ficheiro pickle")
            print("2. Importar dados em ficheiro pickle")
            print("3. Gerar Ficheiro Word")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                guardar_ficheiro()
            
            elif opção == 2:
                abrir_ficheiro()
            
            elif opção == 3:
                gerar_word()
                
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")       


while True:
        print("Benvindo(a)!")
        print("1. Criar lista de objectos e calcular Energia Cinética, Potencial Gravitica e Mecânica")
        print("2. Gestão e Edição de objectos")
        print("3. Cálculos adicionais sobre dados registados (Média/Somatórios/Pólos - Maior e Menor)")
        print("4. Importar ou Guardar dados por ficheiro pickle/Word")
        print("5. Gráficos de linhas e de Barras - Relações entre dados e resultados")
        print("6. Tabela de Dados - Data Frame")
        print("7. Saír")
        try:
            escolha = int(input("\nIntroduza a opção: "))
        
            if escolha == 1:
                construir_array()  
                
            elif escolha == 2:
                gestao_array()  
        
            elif escolha == 3:
                calc_adicionais() 
            
            elif escolha == 4:
                armazenamento()  
                
            elif escolha == 5:
                menu_graph()
                
            elif escolha == 6:
                menu_df()
        
            elif escolha == 7:
                print("Obrigado!")
                break
        
            else:
                print("Erro! Escolha errada!")
                
            
        except ValueError:
                print("Erro! Valor deve ser um numero!")   
    