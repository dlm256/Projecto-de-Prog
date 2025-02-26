import pickle
import numpy as np
from docx import Document
import matplotlib.pyplot as plt
import pandas as pd
import mysql.connector as sql

df = pd.DataFrame()


Objectos = []
Massa = []
Velocidade = []
EngCin = []
Altura = []
Epg = []
Em = []



def menu_df():
    while True:
        try:
            print("1. Criar tratamento de dados em Data Frame")
            print("2. Consultar tabela de dados")
            print("3. Calcular estatisticas por coluna")
            print("4. Renomear Coluna")
            print("5. Eliminar coluna da tabela de dados")
            print("6. Editar valor da tabela de dados")
            print("7. Eliminar linha da tabela de dados")
            print("8. Exportar tabela para Excel")
            print("9. Importar tabela de Excel")
            print("10. Gerar Gráfico de Linhas")
            print("11. Gerar Gráfico de Dispersão")
            print("12. Gerar Gráfico Circular")
            print("13. Regressar ao Menu Principal")
            
            opção = int(input("\nIntroduza a opção: " ))
            
            if opção == 1:
                const_DF()
            
            elif opção == 2:
                consultar_DF()
                
            elif opção == 3:
                calcular_estatística()
            
            elif opção == 4:
                renomear_coluna()
            
            elif opção == 5:
                eliminar_coluna()
            
            elif opção == 6:
                editar_df_linha()
                
            elif opção == 7:
                remover_df_linha()
                
            elif opção == 8:
                export_xl()
            
            elif opção == 9:
                importar_xl()
                
            elif opção == 10:
                gráfico_de_linha()
                
            elif opção == 11:
                gráfico_dispersão()
            
            elif opção == 12:
                gráfico_circular_legenda()
                
            elif opção == 13:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print("Erro! Indique um número!") 
                
def validar_colunas(colunas):
    for col in colunas:
        if col not in df.columns:
            print(f"Erro! A coluna {col} não consta do Data Frame! ")
            return False
        return True
    
    

def gráfico_de_linha():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    print("Colunas no Data Frame: ", list(df.columns))
    x_col = input("Indique o eixo de x: ")
    y_col = input("Indique o eixo de y: ")
    if validar_colunas([x_col,y_col]):
        df.plot(kind='line', x=x_col, y=y_col, legend=True)
        plt.title("Gráfico de Linha")
        legenda = input("Indique a legenda do gráfico: ")
        plt.legend(title=legenda)
        plt.show()
        
def gráfico_dispersão():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    print("Colunas no Data Frame: ", list(df.columns))
    x_col = input("Indique o eixo de x: ")
    y_col = input("Indique o eixo de y: ")
    if validar_colunas([x_col,y_col]):
        df.plot(kind='scatter', x=x_col, y=y_col, legend=True)
        plt.title("Gráfico de Dispersão")
        plt.show()
        
def gráfico_circular_legenda():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    print("Colunas no Data Frame: ", list(df.columns))
    col = input("Indique a coluna para os dados do gráfico circular: ")
    legenda_col = input("Indique a coluna para definir a legenda dos valores: ")
    if validar_colunas([col, legenda_col]):
        df[col].value_counts().plot(kind='pie', autopct='%1.1f%%', labels = df[legenda_col], legend=True)
        plt.title("Gráfico Circular")
        plt.show()
                

def renomear_coluna():
    global df
    if df.empty:
        print("Data Frame vazio! ")
        return
    while True:  
        
        nome_coluna = input("Indique o nome da Coluna a alterar o nome: ")
           
        if nome_coluna not in df.columns:
            print(f"Erro! A Coluna '{nome_coluna}' não consta no Data Frame! ")
        else:     
            break
        
    while True:
        novo_nome_coluna = input("Introduza o novo nome da coluna: ").strip()
        if not novo_nome_coluna:
            print("Erro, o nome da coluna não pode ser vazio! ")
        elif novo_nome_coluna in df.columns:
            print(f"Erro! A coluna '{novo_nome_coluna}' já consta no Data Frame! ")
        else:
            break
    
    df.rename(columns={nome_coluna: novo_nome_coluna}, inplace=True)
    print(f"A coluna '{nome_coluna}' foi renomeada com sucesso para '{novo_nome_coluna}'")
    print()
                
def calcular_estatística():
    global df
    if df.empty:
        print("Data Frame vazio! ")
        return
    while True:
        while True:
            colunas_numericas = df.select_dtypes(include=['float'])
            print("Colunas numéricas disponíveis: ", ", ".join(colunas_numericas.columns))
            
            coluna = input("Indique em que coluna pretende fazer o cálculo: ")
            break
        if coluna not in colunas_numericas.columns:
            print(f"Erro! A coluna '{coluna}' não consta desta tabela! ")
        else:
            break
        
    
    while True:
        print("Selecione o cálculo a efetuar: ")
        print("1. Média ")
        print("2. Somatório ")
        print("3. Valor Maior ")
        print("4. Valor menor ")
        print("5. Mediana ")
        print("6. Retornar ao menu anterior ")
        
        escolha = int(input("Indique o cálculo: "))
        
        if escolha == 1:
            print(f"Média dos valores da coluna {coluna}: {df[coluna].mean()}")
        elif escolha == 2:
            print(f"Somatório dos valores da coluna {coluna}: {df[coluna].sum()}")    
        elif escolha == 3:
            print(f"O valor mais alto da coluna {coluna} é: {df[coluna].max()}") 
        elif escolha == 4:  
            print(f"O valor mais baixo da coluna {coluna} é: {df[coluna].min()}")   
        elif escolha == 5:
            print(f"A mediana dos valores da coluna {coluna} é: {df[coluna].median()}")
        elif escolha == 6:
            print("A voltar ao menu anterior...")
            break
        else:
            print("Erro! A opção é inválida! ")    
                                
        print()        


def editar_df_linha():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    print(df)
    while True:
        while True:
            try:
                linha = int(input("Introduza o nr da linha a alterar: "))
                break
            except ValueError:
                print("Erro! Indique um número!") 
                
            if linha not in df.index:
                print(f"Erro! O índice {linha} não consta na tabela!")
            else:
                break
        while True:
            coluna = input("Introduza a coluna a alterar: ")       
            break
            
        if coluna not in df.columns:
            print(f"Erro! A coluna {coluna} não consta na tabela!")
        else:
            break
        
        
        
    novo_dado = input("Introduza um novo valor: ")
    df.at[linha,coluna] = novo_dado
    print("Dado alterado com sucesso!")
            

def remover_df_linha():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    print(df)
    while True:
        while True:
            try:
                linha = int(input("Introduza o nr da linha do valor a eliminar: "))
                break
            except ValueError:
                print(" Erro! Indique um número!") 
                
        if linha not in df.index:
            print(f"Erro! O índice {linha} não consta na tabela!")
        else:
            break
    
    df.drop(index = linha, inplace=True)
    print("Linha removida com sucesso!")


def const_DF():
    global df
    if not Objectos:
        print("A lista está vazia!")
    else:
        df = pd.DataFrame({"Objecto": Objectos, "Massa(kg)": Massa, "Velocidade(m/s)": Velocidade, "Altura(m)": Altura, "Energia Cinética(J)": EngCin, "Energia Potencial Gravítica(J)": Epg, "Energia Mecânica(J)":Em})
        print("\nData Frame criado com sucesso!\n") 
        return df

def consultar_DF():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    
    print("Composição do Data Frame: \n")
    print(df)
    print()
    
def eliminar_coluna():
    global df
    if df.empty:
        print("Data Frame vazio!\n")
        return
    
    nome_coluna = input("Indique o nome da Coluna a eliminar: \n") 
    
    if nome_coluna not in df.columns:
        print(f"Erro! A coluna '{nome_coluna}' não consta no Data Frame!\n")
        return None  
    
    df.drop(columns=[nome_coluna], inplace=True)
    print(f"Coluna '{nome_coluna}' eliminada com sucesso!\n")
    print()
    
def export_xl():
    global df
    if df.empty:
        print("Data Frame vazio! ")
        return
    nome_ficheiro = input("Introduza o nome do ficheiro a guardar em excel: ")
    
    if not nome_ficheiro.endswith('.xlsx'):
        nome_ficheiro += '.xlsx'
        
    try:
        df.to_excel(nome_ficheiro, index=False)
        print(f"Dados do Data Frame exportados para: {nome_ficheiro}\n")
    except Exception as e:
        print(f"Erro a gerar ficheiro: {e} \n")
    print()

def importar_xl():
    global df
    
    nome_ficheiro = input("Introduza a localização do ficheiro a importar: \n")
    
    if not nome_ficheiro.endswith('.xlsx'):
        nome_ficheiro += '.xlsx'
        
    try:
        df = pd.read_excel(nome_ficheiro)
        print(f"{nome_ficheiro} importado com sucesso!")
    except FileNotFoundError:
        print("Erro! Ficheiro não encontrado!")
    except Exception as e:
        print(f"Erro a gerar ficheiro: {e} \n")
    print()


def construir_array():
    while True:
        try:
            print("1. Registar dados")
            print("2. Regressar ao Menu Principal")
            opção = int(input("\nIntroduza a opção: " ))
            if opção == 1:
                
                objecto = input("Indique o objecto a adicionar: ")
                while True:
                    try:
                        massa = float(input("Indique a massa em kg: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                while True:
                    try:        
                        velocidade = float(input("Indique a velocidade em m/s: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                while True:
                    try: 
                        altura = float(input("Indique a altura do objecto em relação ao nível do mar em metros: "))
                        break
                    except ValueError:
                        print("Erro! O valor deve ser numerico!")
                calc_ec = 0.5*(massa*(velocidade**2))
                calc_epg = massa*altura*9.8
                calc_em = calc_ec + calc_epg 
                Objectos.append(objecto)
                Massa.append(massa)
                Velocidade.append(velocidade)
                Altura.append(altura)
                EngCin.append(calc_ec)
                Epg.append(calc_epg)
                Em.append(calc_em)
                print("Registado!")
                
            elif opção == 2:
                print("Regressar ao menu principal...")
                break
            else:
                print("Opção inválida! Seleccione novamente: ")
                    
        except ValueError:
            print(" Erro! Indique um número!")

    
            
            

def calc_media ():
    if not Objectos:
        print("A lista está vazia!")      
    else:                               
        media_massa = np.mean(Massa)
        media_velocidade = np.mean(Velocidade)
        media_altura = np.mean(Altura)
        media_engcin = np.mean(EngCin)
        media_epg = np.mean(Epg)
        media_em = np.mean(Em)
        print(f"Média das massas: {media_massa} kg") 
        print(f"Média das velocidades: {media_velocidade} m/s")
        print(f"Média das alturas: {media_altura} m")
        print(f"Média de Energia Cinética: {media_engcin} joules")
        print(f"Média da Energia potencial Gravítica: {media_epg} joules")
        print(f"Média da Energia Mecânica: {media_em} joules")
        
def somatório():
    if not Objectos:
        print("A lista está vazia!") 
    else:
        sum_massa = sum(Massa)
        sum_vel = sum(Velocidade)
        sum_h = sum(Altura)
        sum_epg = sum(Epg)
        sum_EngCin = sum(EngCin)
        sum_em = sum(Em)
        print(f"Somatório massas: {sum_massa} kg")
        print(f"Somatório velocidades: {sum_vel}")
        print(f"Somatório velocidades: {sum_h}")
        print(f"Somatório Energia Cinética: {sum_EngCin}")
        print(f"Somatório Energia potencial Gravítica: {sum_epg}")
        print(f"Somatório Energia Mecânica: {sum_em}")
        
def maior_menor():
    if not Objectos:
        print("A lista está vazia!") 
    else:
        maior_valor_massa = np.max(Massa)
        menor_val_massa = np.min(Massa)
        print(f"Massa \nMaior valor: {maior_valor_massa} kg \nMenor valor {menor_val_massa} kg")

        maior_valor_vel = np.max(Velocidade)
        menor_val_vel = np.min(Velocidade)
        print(f"Velocidade \nMaior valor: {maior_valor_vel} m/s \nMenor valor {menor_val_vel} m/s")
        
        maior_valor_h = np.max(Altura)
        menor_val_h = np.min(Altura)
        print(f"Altura \nMaior valor: {maior_valor_h} m \nMenor valor {menor_val_h} m")
        
        maior_valor_engcin = np.max(EngCin)
        menor_val_engcin = np.min(EngCin)
        print(f"Energia Cinética \nMaior valor: {maior_valor_engcin} joules \nMenor valor {menor_val_engcin} joules")
        
        maior_valor_epg = np.max(Epg)
        menor_val_epg = np.min(Epg)
        print(f"Energia Potencial Gravítica \nMaior valor: {maior_valor_epg} joules \nMenor valor {menor_val_epg} joules")
        
        maior_valor_em = np.max(Em)
        menor_val_em = np.min(Em)
        print(f"Energia Mecânica \nMaior valor: {maior_valor_em} joules \nMenor valor {menor_val_em} joules")

   
def consultar_array():
    if not Objectos:
        print("A lista está vazia!")
    else:
        print("Lista de Objectos:\n")
        for i, (objecto, massa, velocidade, altura, calc_ec, calc_epg, calc_em) in enumerate(zip(Objectos,Massa,Velocidade,Altura,EngCin,Epg,Em), start=1):
            print(f"Objecto {i}: {objecto}")
            print(f"Massa: {massa} kg")
            print(f"Velocidade: {velocidade} m/s")
            print(f"Altura: {altura} m")
            print(f"Energia Cinética: {calc_ec} joules")
            print(f"Energia Potencial Gravítica: {calc_epg} joules")
            print(f"Energia Mecânica: {calc_em} joules\n")
            
def eliminar_objectos():
    if not Objectos:
        print("A lista está vazia!")
    else:
        consultar_array()
        try:
            while True:
                id = int(input("\nIndique o id do objecto a eliminar: "))
                if 1 <= id <= len(Objectos):
                    objecto_eliminado = Objectos.pop(id -1)
                    massa_eliminada = Massa.pop(id -1)
                    velocidade_eliminada = Velocidade.pop(id -1)
                    h_eliminada = Altura.pop(id -1)
                    ec_eliminada = EngCin.pop(id -1)
                    epg_eliminada = Epg.pop(id -1)
                    em_eliminada = Em.pop(id -1)
                    print(f"\nItem com id {id} eliminado!")
                    break
        except ValueError:
            print("Erro! O id deve ser um numero!")
            
def alterar_objectos():
    if not Objectos:
        print("A lista está vazia!")
    
    else:
        consultar_array()
        try:
            while True:
                id = int(input("\nIndique o id do objecto a alterar: "))
                if 1 <= id <= len(Objectos): 
                    objecto_actualizado = input("Indique um novo objecto: ")
                    massa_actualizada = float(input("Indique uma nova massa: "))
                    h_actualizada = float(input("Indique uma nova altura: "))
                    velocidade_actualizada = float(input("Indique uma nova velocidade: "))
                    
                    Objectos[id -1] = objecto_actualizado
                    Massa[id -1] = massa_actualizada
                    Velocidade [id -1] = velocidade_actualizada
                    Altura [id -1] = h_actualizada
                    EngCin[id -1] = 0.5*(massa_actualizada*(velocidade_actualizada**2))
                    Epg[id -1] = massa_actualizada*h_actualizada*9.8
                    Em[id -1] = EngCin[id -1] + Epg[id -1]
                    print("\nObjecto actualizado com sucesso!")
                    break
                else:
                    print("Erro! O id é inválido!")
        except ValueError:
                print("Erro! O id deve ser um numero!")

def guardar_ficheiro ():
    if not Objectos:
        print("A lista está vazia!")
    else:
        try:
            file_path = input("Introduza a directoria para guardar, incluíndo o nome final do ficheiro no fim:  ")
            if not file_path.endswith('.pkl'):
                file_path += '.pkl'
            with open(file_path, "wb") as ficheiro:
                pickle.dump((Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em), ficheiro)
                print("Exportado com sucesso!")
        except Exception as e:
            print(f"Erro ao gerar ficheiro: {e}")


def abrir_ficheiro ():
    global Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em
    
    try:
        file_path = input("Introduza a directoria para abrir doc: ")
        if not file_path.endswith('.pkl'):
                file_path += '.pkl'
        with open(file_path, "rb") as ficheiro:
            Objectos, Massa, Altura, Velocidade, EngCin, Epg, Em  = pickle.load(ficheiro)
            print("Ficheiro importado!")
    except FileNotFoundError:
        print("Erro! Ficheiro não encontrado!")
    except Exception as e:
        print(f"Erro a abrir ficheiro: {e}")
        
def gerar_word():
    if not Objectos:
        print("A lista está vazia!")
    else:
        try:
            doc = Document()
            título = input("Indique titulo: ")
            tamanho = int(input("Indique o tamanho de letra em número inteiro(0 é o tamanho de letra maior possível):"))
            doc.add_heading(título, tamanho)

            for i, (objecto, massa, altura, velocidade, calc_ec, calc_epg, calc_em) in enumerate(zip(Objectos,Massa,Altura,Velocidade,EngCin,Epg,Em), start=1):
                doc.add_paragraph(f"{i} - {objecto} com {massa} kg e {velocidade} m/s produz {calc_ec} joules (Energia Cinética), {calc_epg} joules para {altura} metros (Energia Potencial Gravitica) e totalizando {calc_em} joules (Energia Mecânica) \n")
                
            file_path = input("Introduza a directoria para guardar, incluíndo o nome final do ficheiro no fim:  ")
            if not file_path.endswith('.docx'):
                file_path += '.docx'
            doc.save(file_path)
            print("Doc Word gerado com sucesso!")
        except Exception as e:
            print(f"Erro ao gerar ficheiro: {e}")

def graph():
    while True:
        try:
            print("1. Gráfico de Linha da Relação Velocidade/Energia cinética")
            print("2. Gráfico de Linha da Relação Massa/Energia cinética")
            print("3. Gráfico de Linha da Relação Massa/Energia Potencial Gravítica")
            print("4. Gráfico de Linha da Relação Altura/Energia Potencial Gravítica")
            print("5. Gráfico de Linha da Relação Velocidade/Energia cinética")
            print("6. Gráfico de Linha da Relação Energia Mecânica/Energia Potencial Gravítica")
            print("7. Gráfico de Linha da Relação Energia Mecânica/Energia Cinética")
            print("8. Saír")
            
            opção = int(input("Introduza a opção: " ))
            if opção == 1: 
                plt.plot(Velocidade,EngCin)
                plt.xlabel('Velocidade (m/s)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Velocidade/Energia cinética')
                plt.show()
            elif opção == 2:
                plt.plot(Massa,EngCin)
                plt.xlabel('Massa (kg)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Massa/Energia cinética')
                plt.show()
            elif opção == 3:
                plt.plot(Massa,Epg)
                plt.xlabel('Massa (kg)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Massa/Energia Potencial Gravítica')
                plt.show()
            elif opção == 4:    
                plt.plot(Altura,Epg)
                plt.xlabel('Altura (m)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Altura/Energia Potencial Gravítica')
                plt.show()
            elif opção == 5:   
                plt.plot(EngCin,Epg)
                plt.xlabel('Energia Cinética (joules)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Velocidade/Energia cinética')
                plt.show()
            elif opção == 6:    
                plt.plot(Em,Epg)
                plt.xlabel('Energia Mecânica (joules)')
                plt.ylabel('Energia Potencial Gravítica (joules)')
                plt.title('Gráfico de Linha da Relação Energia Mecânica/Energia Potencial Gravítica')
                plt.show()
            elif opção == 7:    
                plt.plot(Em,EngCin)
                plt.xlabel('Energia Mecânica (joules)')
                plt.ylabel('Energia Cinética (joules)')
                plt.title('Gráfico de Linha da Relação Energia Mecânica/Energia Cinética')
                plt.show()
            elif opção == 8:
                print("A saír para menu anterior...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
        except ValueError:
                print(" Erro! Indique um número!")
    
def graph_bar():
    while True:
        try:
            print("1. Gráfico de Barras - Velocidade")
            print("2. Gráfico de Barras - Massa")
            print("3. Gráfico de Barras - Altura")
            print("4. Gráfico de Barras - Energia Cinética")
            print("5. Gráfico de Barras - Energia Potencial Gravítica")
            print("6. Gráfico de Barras - Energia Mecânica")
            print("7. Saír")
            opção = int(input("Introduza a opção: " ))
            if opção == 1:
                plt.bar(Objectos,Velocidade)
                plt.xlabel('Objecto')
                plt.ylabel('Velocidade(m/s)')
                plt.title('Gráfico de Barras - Velocidade')
                plt.show()
            elif opção == 2:
                plt.bar(Objectos,Massa)
                plt.xlabel('Objecto')
                plt.ylabel('Massa(kg)')
                plt.title('Gráfico de Barras - Massa')
                plt.show()
            elif opção == 3:  
                plt.bar(Objectos,Altura)
                plt.xlabel('Objecto')
                plt.ylabel('Altura (m)')
                plt.title('Gráfico de Barras - Altura')
                plt.show()
            elif opção == 4:
                plt.bar(Objectos,EngCin)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Cinética(joules)')
                plt.title('Gráfico de Barras - Energia Cinética')
                plt.show()
            elif opção == 5:
                plt.bar(Objectos,Epg)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Potencial Gravítica(joules)')
                plt.title('Gráfico de Barras - Energia Potencial Gravítica')
                plt.show()
            elif opção == 6:
                plt.bar(Objectos,Em)
                plt.xlabel('Objecto')
                plt.ylabel('Energia Mecânica(joules)')
                plt.title('Gráfico de Barras - Energia Mecânica')
                plt.show()
            elif opção == 7:
                print("A saír para Menu anterior...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
            
            
        except ValueError:
                print(" Erro! Indique um número!")
            
def menu_graph():
    while True:
        try:
            print("1. Gráficos de linhas da relação entre os valores")
            print("2. Gráficos de Barras para atribuição de valor por objecto")
            print("3. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                graph()
            
            elif opção == 2:
                graph_bar()
            
            elif opção == 3:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def calc_adicionais():
    while True:
        try:
            print("1. Calcular a média")
            print("2. Calcular o somatório de cada conjunto de dados")
            print("3. Calcular Maior e Menor de cada conjunto")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                calc_media()
            
            elif opção == 2:
                somatório()
                
            elif opção == 3:
                maior_menor()
            
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def gestao_array():
    while True:
        try:
            print("1. Listar objectos")
            print("2. Eliminar objectos")
            print("3. Alterar objectos")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                consultar_array()
            
            elif opção == 2:
                eliminar_objectos()
                
            elif opção == 3:
                alterar_objectos()
            
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")

def armazenamento():
    while True:
        try:
            print("1. Guardar dados em ficheiro pickle")
            print("2. Importar dados em ficheiro pickle")
            print("3. Gerar Ficheiro Word")
            print("4. Regressar ao Menu Principal")
            
            opção = int(input("Introduza a opção: " ))
            
            if opção == 1:
                guardar_ficheiro()
            
            elif opção == 2:
                abrir_ficheiro()
            
            elif opção == 3:
                gerar_word()
                
            elif opção == 4:
                print("A regressar ao Menu Principal...")
                break
            else:
                print("Erro! Seleccione uma das opções!")
                
        except ValueError:
                print(" Erro! Indique um número!")       


def ligação_DB():
    try:
        ligação = sql.connect(
            host = "localhost",
            user = "DLM256",
            password = "OchedinValannor_256",
            database = "python"
    )
        print("Ligação sucedida!")
        return ligação
    except sql.Error as erro:
        print(f"Erro a estabelecer ligação à BD: {erro}")
        return None
   
def const_DB(cursor):
    cmd_sql = """
    CREATE TABLE IF NOT EXISTS Objectos(
        id INT AUTO_INCREMENT PRIMARY KEY,
        Objecto VARCHAR(255),
        Massa_Kg FLOAT,
        Velocidade_m_s FLOAT,
        Altura_m FLOAT,
        Energia_Cinética_J FLOAT,
        Energia_Potencial_Gravítica_J FLOAT,
        Energia_Mecânica_J FLOAT
       
    )
   
    """
   
    cursor.execute(cmd_sql)
    print("Tabela criada com sucesso!\n")
    connect_DB.commit()
   
def ins_DB(cursor):
    if not Objectos:
        print("Lista vazia!")
        return None
    
    cmd_sql = "INSERT INTO Objectos (Objecto,Massa_Kg,Velocidade_m_s,Altura_m,Energia_Cinética_J,Energia_Potencial_Gravítica_J,Energia_Mecânica_J) VALUES (%s, %s, %s, %s, %s, %s, %s)"

    Arrays = list(zip(Objectos,Massa,Velocidade,Altura,EngCin,Epg,Em))
   
    cursor.executemany(cmd_sql,Arrays)
   
   
    print("Dados adicionados com sucesso! ")
    connect_DB.commit()
 
def search_DB(cursor):
    cmd_sql = "SELECT * FROM Objectos"
    cursor.execute(cmd_sql)
   
    resultados = cursor.fetchall()
    for resultado in resultados:
        print(resultado)
    
    count_DB(cursor)
    avg_DB(cursor)
   
    connect_DB.commit()
    
def erase_DB(cursor):
    search_DB(cursor)
    try:
        id = int(input("Introduza o id do Objectos a remover: "))
        cursor.execute("SELECT * FROM Objectos WHERE id = %s", (id,))
        resultado = cursor.fetchone()
        if resultado:
            cursor.execute("DELETE FROM Objectos WHERE id = %s", (id,))
        else:
            print("Erro! ID não encontrado!")
    except ValueError:
                print("Erro! Valor deve ser um numero!")
    connect_DB.commit
    
def erase_DB_nome(cursor):
    search_DB(cursor)
    
    id = input("Introduza o nome do Objectos a remover: ")
    cursor.execute("SELECT * FROM Objectos WHERE Objecto = %s", (id,))
    resultado = cursor.fetchall()
    if resultado:
        cursor.execute("DELETE FROM Objectos WHERE Objecto = %s", (id,))
        print(f"Linha com Objecto {id} eliminado com sucesso!")
    else:
        print("Erro! Objecto não encontrado!")
    connect_DB.commit
    
def SQL_import(cursor):
    global Objectos, Massa, Velocidade, Altura, EngCin, Epg, Em

    data_fields = {
        "Objecto": Objectos,
        "Massa_Kg": Massa,
        "Velocidade_m_s": Velocidade,
        "Altura_m": Altura,
        "Energia_Cinética_J": EngCin,
        "Energia_Potencial_Gravítica_J": Epg,
        "Energia_Mecânica_J": Em
    }

    for column, array in data_fields.items():
        cursor.execute(f"SELECT {column} FROM Objectos")
        array.extend([row[0] for row in cursor.fetchall()])
        print(array)

    print("Dados importados com sucesso!")

    connect_DB.commit()

    
    
def edit_DB(cursor):
    search_DB(cursor)
    try:
        id = int(input("Introduza o id do Objecto a editar: "))
        cursor.execute("SELECT * FROM Objectos WHERE id = %s", (id,))
        resultado = cursor.fetchone()
        
        
        if resultado:
            while True:
                print("1. Editar nome do Objecto")
                print("2. Editar massa do Objecto")
                print("3. Editar velocidade do Objecto")
                print("4. Editar altura do Objecto")
                print("5. Editar valores da linha na totalidade")
                print("6. Saír")
                try:
                    escolha = int(input("\nSeleccione uma opção: "))
                    if escolha == 1:
                        objecto = input("Indique o objecto a adicionar: ")
                        cursor.execute("UPDATE Objectos SET Objecto = %s WHERE id = %s" , (objecto, id))
                        print("Actualização com sucesso!")
                        
                    
                    elif escolha == 2:
                        while True:
                            try:
                                massa = float(input("Indique a massa em kg: "))
                                cursor.execute("UPDATE Objectos SET Massa_Kg = %s WHERE id = %s" , (massa, id))
                                print("Actualização com sucesso!")
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                                
                    elif escolha == 3:
                        while True:
                            try:        
                                velocidade = float(input("Indique a velocidade em m/s: "))
                                cursor.execute("UPDATE Objectos SET Velocidade_m_s = %s WHERE id = %s" , (velocidade, id))
                                print("Actualização com sucesso!")
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                                
                    elif escolha == 4:
                        while True:
                            try: 
                                altura = float(input("Indique a altura do objecto em relação ao nível do mar em metros: "))
                                cursor.execute("UPDATE Objectos SET Altura_m = %s WHERE id = %s" , (altura, id))
                                print("Actualização com sucesso!")
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                                
                    elif escolha == 5:
                        objecto = input("Indique o objecto a adicionar: ")
                        while True:
                            try:
                                massa = float(input("Indique a massa em kg: "))
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                        while True:
                            try:        
                                velocidade = float(input("Indique a velocidade em m/s: "))
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                        while True:
                            try: 
                                altura = float(input("Indique a altura do objecto em relação ao nível do mar em metros: "))
                                break
                            except ValueError:
                                print("Erro! O valor deve ser numerico!")
                        cursor.execute("UPDATE Objectos SET Objecto = %s, Massa_Kg = %s, Velocidade_m_s = %s, Altura_m = %s WHERE id = %s" , (objecto, massa, velocidade, altura, id))
                        print("Actualização com sucesso!")
                    
                    elif escolha == 6:
                        print("A retornar ao menu anterior...")
                        break
                    
                    else:
                        print("Escolha errada!")
                        
                except ValueError:
                    print("Erro! Valor deve ser um numero!")
                    
            
        else:
            print("Erro! ID não encontrado!")
    except ValueError:
                print("Erro! Valor deve ser um numero!")
                
    connect_DB.commit()
    
def act_resultados_SQL(cursor):
    search_DB(cursor)
    while True:
        try:
            print("Actualização de resultados para novos dados introduzidos: ")
            print("\n1.Actualizar resultados por ID")
            print("2.Retornar ao menu anterior")
            escolha = int(input("\nIntroduza uma opção: "))
            if escolha == 1:
                try:
                    id = int(input("\nIntroduza o id para actualizar os resultados: "))
                    cursor.execute( "SELECT Massa_Kg, Velocidade_m_s, Altura_m FROM Objectos WHERE id = %s", (id,))
                    resultado = cursor.fetchone()

                    if resultado:
                        massa, velocidade, altura = resultado

                        if massa is not None and velocidade is not None and altura is not None:
                            
                            calc_ec = 0.5 * massa * (velocidade ** 2)
                            calc_epg = massa * altura * 9.8
                            calc_em = calc_ec + calc_epg

                            cursor.execute( "UPDATE Objectos SET Energia_Cinética_J = %s,Energia_Potencial_Gravítica_J = %s,Energia_Mecânica_J = %s WHERE id = %s",(calc_ec, calc_epg, calc_em, id))
                            print("Actualização com sucesso!")
                        else:
                            print("Erro! Valor inválido!")

                    else:
                        print("Erro! ID não encontrado!")

                except ValueError:
                    print("Erro! O valor do ID deve ser um número.")
            
            elif escolha == 2:
                print("A retornar ao menu anterior...\n")
                break
            else:
                print("Erro! Escolha errada!")
                
        except ValueError:
                    print("Erro! O valor do ID deve ser um número.")

    connect_DB.commit()
    
def count_DB(cursor):
    cursor.execute("SELECT COUNT(*) FROM Objectos")
    total = cursor.fetchone()[0]
    print(f"\nO total de Objectos é {total}.")
    connect_DB.commit()
    
def avg_DB(cursor):
    cursor.execute( "SELECT AVG(Massa_Kg),AVG(Velocidade_m_s),AVG(Altura_m),AVG(Energia_Cinética_J),AVG(Energia_Potencial_Gravítica_J),AVG(Energia_Mecânica_J)FROM Objectos")
    media = cursor.fetchone()
    
    print(f"\nMédia da Massa: {media[0]:.2f} kg")
    print(f"Média da Velocidade: {media[1]:.2f} m/s")
    print(f"Média da Altura: {media[2]:.2f} m")
    print(f"Média da Energia Cinética: {media[3]:.2f} J")
    print(f"Média da Energia Potencial Gravítica: {media[4]:.2f} J")
    print(f"Média da Energia Mecânica: {media[5]:.2f} J")
    connect_DB.commit()

def verificar_tabela(cursor, nome_tabela):
    cursor.execute("SHOW TABLES LIKE %s", (nome_tabela,))
    return cursor.fetchone()    
    
def filtrar_dados(cursor):
    nome_tabela = input("Indique o nome da tabela: ").strip()

    if not verificar_tabela(cursor, nome_tabela):
        print("Erro! A tabela não existe!")
        return

    # Get column names and their types
    cursor.execute(f"DESCRIBE {nome_tabela}")
    colunas_info = cursor.fetchall()

    colunas = {col[0]: col[1] for col in colunas_info}  # Dictionary {column_name: column_type}
    
    print("Colunas disponíveis:", ", ".join(colunas.keys()))
    coluna = input("Indique a coluna para filtrar: ").strip()

    if coluna not in colunas:
        print("Erro! A coluna não existe na tabela.")
        return

    # Infer data type and convert input accordingly
    tipo_coluna = colunas[coluna].lower()

    valor = input(f"Introduza o valor para filtrar na coluna '{coluna}': ").strip()

    # Convert input value based on column type
    if "int" in tipo_coluna:
        try:
            valor = int(valor)
        except ValueError:
            print("Erro! O valor deverá ser um número inteiro.")
            return
    elif "decimal" in tipo_coluna or "float" in tipo_coluna or "double" in tipo_coluna:
        try:
            valor = float(valor)
        except ValueError:
            print("Erro! O valor deverá ser um número decimal.")
            return
    elif "date" in tipo_coluna or "timestamp" in tipo_coluna:
        from datetime import datetime
        try:
            valor = datetime.strptime(valor, "%Y-%m-%d")  # Adjust format if necessary
        except ValueError:
            print("Erro! A data deve estar no formato YYYY-MM-DD.")
            return
    # No conversion needed for strings (VARCHAR, TEXT, etc.)

    # Execute the filtered query
    try:
        cursor.execute(f"SELECT * FROM {nome_tabela} WHERE {coluna} = %s", (valor,))
        resultados = cursor.fetchall()

        if resultados:
            print(f"\nResultados filtrados por {coluna} = {valor}:")
            for resultado in resultados:
                print(resultado)
        else:
            print("Nenhum registo encontrado neste critério.")

    except sql.Error as erro:
        print(f"Erro ao filtrar os dados: {erro}")

        
def listar_tabelas(cursor):
    try:
        cursor.execute("SHOW TABLES")
        tabelas = cursor.fetchall()
        
        if tabelas:
            print("\nTabelas disponíveis na base de dados:")
            print("\n".join(tabela[0] for tabela in tabelas))
        else:
            print("Nenhuma tabela encontrada na base de dados.")

    except sql.Error as erro:
        print(f"Erro ao listar as tabelas: {erro}")

    
def menu_sql():
    while True:
        print("\nArmazenamento em Base de Dados(SQL)")
        print("1. Criar Base de Dados")
        print("2. Inserir registos em SQL na Base de Dados")
        print("3. Exibir dados armazenados")
        print("4. Eliminar linha por ID")
        print("5. Eliminar linha por Objecto")
        print("6. Importar de SQL")
        print("7. Editar dados por ID")
        print("8. Atualizar Resultados por ID")
        print("9. Listar Tabelas Disponíveis")
        print("10. Filtrar Dados em uma Tabela")
        print("11. Sair")

        try:
            escolha = int(input("\nIntroduza a opção: "))

            if escolha == 1:
                const_DB(cursor)
            elif escolha == 2:
                ins_DB(cursor)
            elif escolha == 3:
                search_DB(cursor)
            elif escolha == 4:
                erase_DB(cursor)
            elif escolha == 5:
                erase_DB_nome(cursor)
            elif escolha == 6:
                SQL_import(cursor)
            elif escolha == 7:
                edit_DB(cursor)
            elif escolha == 8:
                act_resultados_SQL(cursor)
            elif escolha == 9:
                listar_tabelas(cursor)
            elif escolha == 10:
                filtrar_dados(cursor)
            elif escolha == 11:
                print("A retornar ao Menu principal...")
                break
            else:
                print("Erro! Escolha errada!")

        except ValueError:
            print("Erro! Valor deve ser um número!")

 
connect_DB = ligação_DB()
if connect_DB:
    cursor = connect_DB.cursor()
 
 
    while True:
        print("Benvindo(a)!")
        print("1. Criar lista de objectos e calcular Energia Cinética, Potencial Gravitica e Mecânica")
        print("2. Gestão e Edição de objectos")
        print("3. Cálculos adicionais sobre dados registados (Média/Somatórios/Pólos - Maior e Menor)")
        print("4. Importar ou Guardar dados por ficheiro pickle/Word")
        print("5. Gráficos de linhas e de Barras - Relações entre dados e resultados")
        print("6. Tabela de Dados - Data Frame")
        print("7. Armazenamento em Base de Dados(SQL)")
        print("8. Saír")
        try:
            escolha = int(input("\nIntroduza a opção: "))
       
            if escolha == 1:
                construir_array()  
               
            elif escolha == 2:
                gestao_array()  
       
            elif escolha == 3:
                calc_adicionais()
           
            elif escolha == 4:
                armazenamento()  
               
            elif escolha == 5:
                menu_graph()
               
            elif escolha == 6:
                menu_df()
           
            elif escolha == 7:
                menu_sql()
       
            elif escolha == 8:
                print("Obrigado!")
                break
       
            else:
                print("Erro! Escolha errada!")
               
            connect_DB.commit()
             
           
        except ValueError:
                print("Erro! Valor deve ser um numero!")
                 
    cursor.close()
    connect_DB.close()              
else:
    print("Erro de ligação!")
